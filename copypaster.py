import json
from typing import Any
from docx import Document
from docx.shared import Inches
from PIL import Image
from config import workdir, src_img_prefix, sol_img_ext


with open(f"{workdir}/problem_ranges_edited.json", "r") as f:
    problem_ranges: dict[int, Any] = json.load(f)

doc = Document()
for i, A in problem_ranges.items():
    src_img_path = f"{workdir}/src/{src_img_prefix}{str(i).zfill(2)}.png"
    doc.add_picture(src_img_path, width = Inches(6.0))
    for n in range(A[0], A[1] + 1):
        zfilled_number = str(n).zfill(3)
        sol_path = f"{workdir}/sol/{zfilled_number}.{sol_img_ext}"
        # Open an image file
        with Image.open(sol_path) as img:
            # Save it back in the same format
            tmp_path = f"./Temp/{zfilled_number}.{sol_img_ext}"
            img.save(tmp_path)
        doc.add_picture(tmp_path, width = Inches(6.0))
        print("#", end = "")
    print()
doc.save(f"{workdir}/output.docx")