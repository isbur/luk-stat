import json
import cv2
from docx import Document
from docx.shared import Inches
from PIL import Image
from visual import PSM3
from config import workdir, number


print("Processing images...")
# images = []
doc = Document()
problem_ranges: dict[int, tuple[int, int]] = {}
# for i in range(1, 2):
for i in range(1, number + 1):
    src_img_path = f"{workdir}/src/input-{str(i).zfill(2)}.png"
    img = cv2.imread(src_img_path)

    with open(f"{workdir}/psm3/{str(i).zfill(2)}.json", "r") as f:
        d: dict[str, list[int|float|str]] = json.load(f)
    
    problems = PSM3(img, d)
    # Упрощённый подход
    numbers = [p.number for p in problems]
    problem_ranges[i] = (min(numbers), max(numbers))
    # if problems.check_sequence():
    #     numbers = [p.number for p in problems]
    #     problem_ranges[i] = (min(numbers), max(numbers))
    # else:
    #     problem_ranges[i] = (0, 0)


    # doc.add_picture(src_img_path, width = Inches(6.0))

    # print(problems)

    # for p in problems:
        # https://stackoverflow.com/questions/9084609/how-to-copy-a-image-region-using-opencv-in-python
        # if p.Rect is None:
        #     raise Exception("Problem Rect not initialized")
        # if p.extra:
        #     continue
        # p.statement_img = img[p.Rect.y:p.Rect.v, p.Rect.x:p.Rect.u]
        # zfilled_number = str(p.number).zfill(3)
        # tmp_path = f"./Temp/{zfilled_number}.png"
        # cv2.imwrite(tmp_path, p.statement_img)
        # doc.add_picture(tmp_path, width = Inches(6.0))
        # sol_path = f"./solutions/{zfilled_number}.png"
        # doc.add_picture(sol_path, width = Inches(6.0))

    # images.append(Image.open(f"./Temp/{str(i).zfill(2)}.jpg"))
    
    print("#", end = "")

print()
json_object = json.dumps(problem_ranges, indent=4)
with open(f"{workdir}/problem_ranges.json", "w") as outfile:
    outfile.write(json_object)
# doc.save("output.docx")
# images[0].save(
#     "output.pdf", "PDF", resolution = 100.0, save_all=True, append_images = images[1:]
# )
